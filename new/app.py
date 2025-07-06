# notebook-backend/app.py

from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
import json
import time
from datetime import datetime
from dotenv import load_dotenv
import re

# DeepSeek 相关的导入
from openai import OpenAI
from openai import APIConnectionError, RateLimitError, APIStatusError

# 用于文档内容提取和生成
import docx
from io import BytesIO
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter

# 导入中文分词库 jieba
import jieba

# 导入业务逻辑模块
from llm_interface import DeepSeekLLM
from data_manager import TemplateMethodManager
from text_summarizer import TextSummarizer
from question_rewriter import QuestionRewriter

# 加载 .env 文件中的环境变量
load_dotenv()

# --- Flask 应用配置 ---
app = Flask(__name__)
CORS(app)

# --- 文件存储配置 ---
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    print(f"Created uploads directory at: {os.path.abspath(UPLOAD_FOLDER)}")

# --- DeepSeek API 配置 ---
deepseek_api_key = "sk-945bab02da964bcdaca673485c32dfff"
if not deepseek_api_key:
    print("WARNING: DEEPSEEK_API_KEY not found in .env file or environment variables.")
    print("AI model calls might fail. Please add DEEPSEEK_API_KEY=your_key_here to your .env file.")

DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"
# 直接用于RAG和普通聊天的OpenAI客户端实例（因为RAG Prompting复杂）
openai_client_for_chat_rag = OpenAI(
    api_key=deepseek_api_key,
    base_url=DEEPSEEK_BASE_URL
)
DEFAULT_DEEPSEEK_MODEL = "deepseek-chat"

# --- 新功能模块初始化 ---
llm_client_for_tasks = DeepSeekLLM(api_key=deepseek_api_key, model=DEFAULT_DEEPSEEK_MODEL, base_url=DEEPSEEK_BASE_URL)
TEMPLATES_FILE = "templates_methods.json"
template_manager = TemplateMethodManager(file_path=TEMPLATES_FILE)
text_summarizer_module = TextSummarizer(llm_client=llm_client_for_tasks)
question_rewriter_module = QuestionRewriter(llm_client=llm_client_for_tasks, template_manager=template_manager)

# --- 模拟数据持久化 ---
mock_files_metadata = []
mock_chat_history = []
mock_chat_messages = {}

MOCK_DATA_FILE = 'mock_data.json'


def load_mock_data():
    global mock_files_metadata, mock_chat_history, mock_chat_messages
    if os.path.exists(MOCK_DATA_FILE):
        with open(MOCK_DATA_FILE, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                mock_files_metadata = data.get('files', [])
                mock_chat_history = data.get('chat_history', [])
                mock_chat_messages = data.get('chat_messages', {})
                print(f"Mock data loaded from {MOCK_DATA_FILE}")
            except json.JSONDecodeError as e:
                print(f"Error decoding {MOCK_DATA_FILE}: {e}. Starting with empty data.")
                if os.path.exists(MOCK_DATA_FILE):
                    os.rename(MOCK_DATA_FILE, f"{MOCK_DATA_FILE}.bak_{int(time.time())}")
    else:
        print(f"No {MOCK_DATA_FILE} found, starting with empty data.")


def save_mock_data():
    data = {
        'files': mock_files_metadata,
        'chat_history': mock_chat_history,
        'chat_messages': mock_chat_messages
    }
    with open(MOCK_DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    print(f"Mock data saved to {MOCK_DATA_FILE}")


# --- 文件内容分块函数 ---
def chunktext(text, chunk_size=800, chunk_overlap=100):
    """
    使用 Langchain 的 RecursiveCharacterTextSplitter 进行智能分块。
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=False
    )
    raw_chunks = text_splitter.split_text(text)

    return [c.strip() for c in raw_chunks if c.strip()]


# --- 文件内容提取函数 ---
def extract_text_from_file(filepath, file_type):
    full_text = ""
    try:
        if file_type == 'docx':
            doc = docx.Document(filepath)
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            print(f"Extracted {len(full_text)} characters from DOCX.")
        elif file_type == 'pdf':
            reader = PdfReader(filepath)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
            print(f"Extracted {len(full_text)} characters from PDF using PyPDF2.")
        elif file_type in ['txt', 'md', 'json', 'log', 'py', 'js', 'css', 'html', 'xml']:
            with open(filepath, 'r', encoding='utf-8') as f:
                full_text = f.read()
            print(f"Read {len(full_text)} characters from plain text file.")
        else:
            print(f"Warning: File type {file_type} not supported for text extraction.")
            return [], f"无法提取类型为 {file_type} 的文件内容。"

    except Exception as e:
        print(f"Error extracting text from {filepath} ({file_type}): {e}")
        return [], f"提取文件内容时发生错误: {str(e)}"

    chunks = chunktext(full_text)

    structured_chunks = []
    for i, chunk_text in enumerate(chunks):
        if chunk_text.strip():
            structured_chunks.append({
                "id": f"chunk_{i + 1}",
                "text": chunk_text,
                "original_full_text": full_text
            })
    return structured_chunks, ""


# --- API 路由 ---

# 1. 提供静态文件 (上传的文件)
@app.route('/uploads/<path:filename>')
def uploaded_file_endpoint(filename):
    print(f"\n--- DEBUG: Serving file request ---")
    print(f"Requested filename (from URL): {filename}")
    print(f"Attempting to serve from directory: {os.path.abspath(UPLOAD_FOLDER)}")

    try:
        response = send_from_directory(UPLOAD_FOLDER, filename)
        print(f"Successfully served file: {os.path.join(UPLOAD_FOLDER, filename)}")
        return response
    except Exception as e:
        print(f"ERROR: Failed to serve file '{filename}'. Details: {e}")
        if not os.path.exists(os.path.join(UPLOAD_FOLDER, filename)):
            print(f"  Reason: File does NOT exist at {os.path.join(os.path.abspath(UPLOAD_FOLDER), filename)}")
            return jsonify({'error': 'File not found on server.', 'detail': f'File {filename} does not exist.'}), 404
        else:
            print(f"  Reason: Other error during serving (e.g., permissions).")
            return jsonify({'error': 'Error accessing file.', 'detail': str(e)}), 500


# 2. 获取文件列表
@app.route('/api/files', methods=['GET'])
def get_files():
    current_files_on_disk_names = os.listdir(UPLOAD_FOLDER)
    updated_mock_files_metadata = []

    for filename_on_disk in current_files_on_disk_names:
        filepath = os.path.join(UPLOAD_FOLDER, filename_on_disk)
        if os.path.isfile(filepath):
            file_meta = next((f for f in mock_files_metadata if f['name'] == filename_on_disk), None)
            if not file_meta:
                _, ext = os.path.splitext(filename_on_disk)
                file_type = ext.replace('.', '') if ext else 'unknown'

                file_meta = {
                    'id': f'f{int(time.time() * 1000)}',
                    'name': filename_on_disk,
                    'size': f"{round(os.path.getsize(filepath) / (1024 * 1024), 2)}MB",
                    'type': file_type,
                    'uploadDate': datetime.fromtimestamp(os.path.getmtime(filepath)).strftime('%Y-%m-%d %H:%M'),
                    'file_type_tag': 'unknown'  # 默认值，如果前端没传
                }
            updated_mock_files_metadata.append(file_meta)

    mock_files_metadata[:] = [f for f in updated_mock_files_metadata if
                              os.path.exists(os.path.join(UPLOAD_FOLDER, f['name']))]
    save_mock_data()
    return jsonify(mock_files_metadata)


# 3. 上传文件 (新增接收 file_type_tag)
@app.route('/api/upload', methods=['POST'])
def upload_file():
    print("Received upload request. Files in request.files:")
    for key, value in request.files.items():
        print(f"  Key: {key}, Filename: {value.filename}, Content-Type: {value.content_type}")

    if 'file' not in request.files:
        print("Error: 'file' part not found in request.files.")
        return jsonify({'error': 'No file part in the request. Make sure your form field name is "file".'}), 400

    file = request.files['file']
    file_type_tag = request.form.get('file_type_tag', 'unknown')
    print(f"Received file_type_tag: {file_type_tag}")

    if file.filename == '':
        print("Error: Empty filename.")
        return jsonify({'error': 'No selected file or empty filename.'}), 400

    original_filename = file.filename
    print(f"Original filename from frontend: {original_filename}")

    name_without_ext, ext = os.path.splitext(original_filename)

    cleaned_name_base = re.sub(r'[^\w\u4e00-\u9fa5.\s-]', '', name_without_ext)
    cleaned_name_base = re.sub(r'\s+', '_', cleaned_name_base).strip('_')

    if not cleaned_name_base:
        final_base_name = f"uploaded_file_{int(time.time() * 1000)}"
        print(f"Warning: Cleaned base name was empty, using fallback: {final_base_name}")
    else:
        final_base_name = cleaned_name_base
        print(f"Cleaned base name: {final_base_name}")

    final_filename = final_base_name + ext
    print(f"Final filename for saving: {final_filename}")

    filepath = os.path.join(UPLOAD_FOLDER, final_filename)

    count = 1
    original_filepath = filepath
    while os.path.exists(filepath):
        name_only, ext_only = os.path.splitext(final_filename)
        final_filename = f"{name_only}_{count}{ext_only}"
        filepath = os.path.join(UPLOAD_FOLDER, final_filename)
        count += 1
    if original_filepath != filepath:
        print(f"File already exists, new unique filename: {final_filename}")

    try:
        file.save(filepath)
        print(f"File successfully saved to: {filepath}")

        file_type = ext.replace('.', '') if ext else 'unknown'

        new_file_meta = {
            'id': f'f{int(time.time() * 1000)}',
            'name': final_filename,
            'size': f"{round(os.path.getsize(filepath) / (1024 * 1024), 2)}MB",
            'type': file_type,
            'uploadDate': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'file_type_tag': file_type_tag
        }
        mock_files_metadata.append(new_file_meta)
        save_mock_data()

        return jsonify({
            'message': 'File uploaded successfully',
            'fileName': final_filename,
            'filePath': filepath,
            'fileId': new_file_meta['id'],
            'fileTypeTag': file_type_tag
        }), 200
    except Exception as e:
        print(f"Error saving file: {e}")
        return jsonify({'error': f'Failed to save file: {str(e)}'}), 500


# 4. 获取对话历史列表 (更新以返回 related_file_ids 和 meta)
@app.route('/api/chat-history', methods=['GET'])
def get_chat_history():
    history_with_files = []
    for chat_entry in mock_chat_history:
        entry_copy = chat_entry.copy()
        entry_copy['related_files_meta'] = []
        for file_id in entry_copy.get('related_file_ids', []):
            file_meta = next((f for f in mock_files_metadata if f['id'] == file_id), None)
            if file_meta:
                entry_copy['related_files_meta'].append(file_meta)
        history_with_files.append(entry_copy)
    return jsonify(history_with_files)


# 5. 获取某个对话的消息 (更新以加载相关文件信息)
@app.route('/api/chat/<chat_id>/messages', methods=['GET'])
def get_chat_messages(chat_id):
    messages = mock_chat_messages.get(chat_id, [])

    related_files_meta = []
    chat_entry = next((c for c in mock_chat_history if c['id'] == chat_id), None)
    if chat_entry and chat_entry.get('related_file_ids'):
        for file_id in chat_entry['related_file_ids']:
            file_meta = next((f for f in mock_files_metadata if f['id'] == file_id), None)
            if file_meta:
                related_files_meta.append(file_meta)

    return jsonify({"messages": messages, "related_files_meta": related_files_meta})


# 6. 发送新消息到大模型 (RAG逻辑)
@app.route('/api/chat', methods=['POST'])
def send_chat_message():
    data = request.get_json()
    user_message_content = data.get('message')
    chat_id = data.get('chatId')
    model_settings = data.get('modelSettings', {})
    related_file_ids = data.get('relatedFileIds', [])

    if not user_message_content:
        return jsonify({'error': 'Message content cannot be empty'}), 400

    new_chat_id = chat_id
    current_chat_entry = None

    if not new_chat_id:
        new_chat_id = f'c{int(time.time() * 1000)}'
        new_chat_title = f"新对话 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        current_chat_entry = {
            'id': new_chat_id,
            'title': new_chat_title,
            'lastActive': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'related_file_ids': related_file_ids
        }
        mock_chat_history.insert(0, current_chat_entry)
        mock_chat_messages[new_chat_id] = []
    else:
        for chat_entry in mock_chat_history:
            if chat_entry['id'] == new_chat_id:
                current_chat_entry = chat_entry
                current_chat_entry['related_file_ids'] = related_file_ids
                break

    current_chat_msgs = mock_chat_messages.get(new_chat_id, [])
    current_chat_msgs.append({
        'id': f'm_user_{int(time.time() * 1000)}',
        'sender': 'user',
        'content': user_message_content,
        'timestamp': datetime.now().strftime('%H:%M')
    })
    mock_chat_messages[new_chat_id] = current_chat_msgs

    final_ai_response_json = {"text": "抱歉，AI助手暂时无法响应。请确保您的API Key正确且网络畅通。", "citations": []}

    try:
        messages_for_llm = []

        context_str = ""
        citations_for_response = []

        # --- 1. 文件内容提取和分块 ---
        all_selected_file_chunks = []
        if related_file_ids:
            print(f"DEBUG RAG: Retrieving content from selected files for RAG: {related_file_ids}")
            for file_id in related_file_ids:
                file_meta = next((f for f in mock_files_metadata if f['id'] == file_id), None)
                if file_meta:
                    file_path = os.path.join(UPLOAD_FOLDER, file_meta['name'])
                    file_type = file_meta['type']

                    if os.path.exists(file_path):
                        chunks_from_file, extract_error = extract_text_from_file(file_path, file_type)
                        if chunks_from_file:
                            for chunk in chunks_from_file:
                                chunk['doc_id'] = file_id
                                chunk['doc_name'] = file_meta['name']
                                all_selected_file_chunks.append(chunk)
                        else:
                            print(
                                f"WARNING RAG: Failed to extract chunks from file {file_meta['name']}: {extract_error}")
                    else:
                        print(f"WARNING RAG: File {file_meta['name']} not found on disk for RAG.")
                else:
                    print(f"WARNING RAG: File metadata for ID {file_id} not found for RAG.")

        print(f"DEBUG RAG: Total chunks extracted from selected files: {len(all_selected_file_chunks)}")
        for i, chunk in enumerate(all_selected_file_chunks[:min(3, len(all_selected_file_chunks))]):
            print(
                f"DEBUG RAG: Sample Extracted Chunk {i + 1} (from {chunk.get('doc_name', 'N/A')}, id: {chunk.get('id', 'N/A')}): {chunk['text'][:150]}...")

        # --- 2. 检索最相关的文本块 (使用 jieba 进行中文分词和匹配) ---
        relevant_chunks_for_llm = []
        if all_selected_file_chunks:
            query_segmented_words = jieba.lcut(user_message_content.lower())

            stop_words = {"的", "了", "是", "在", "我", "你", "他", "她", "它", "分析", "文件", "请", "如何", "什么",
                          "这个", "那个", "根据", "文档", "回答", "问题", "并", "和", "或", "等", "以", "从", "中",
                          "于", "对", "为", "所", "其", "则", "将", "与", "关于", "有", "也", "都", "还", "是什么",
                          "的", "地", "得"}
            query_words = {word for word in query_segmented_words if len(word) > 1 and word not in stop_words}

            print(f"DEBUG RAG: Processed Query Words (jieba): {query_words}")

            scored_chunks = []
            for chunk in all_selected_file_chunks:
                if not chunk.get('text'):
                    continue

                chunk_segmented_words = set(jieba.lcut(chunk['text'].lower()))

                score = sum(1 for word in query_words if word in chunk_segmented_words)

                if score > 0:
                    scored_chunks.append((score, chunk))

            print(
                f"DEBUG RAG: Scored Chunks (before sort/limit, jieba): {[(s, c['doc_name'], c['id']) for s, c in scored_chunks[:min(5, len(scored_chunks))]]}...")

            scored_chunks.sort(key=lambda x: x[0], reverse=True)

            MAX_CONTEXT_CHUNKS = 8
            citation_id_counter = 1
            seen_chunk_identifiers = set()

            for score, chunk in scored_chunks:
                if citation_id_counter > MAX_CONTEXT_CHUNKS:
                    break

                chunk_identifier = f"{chunk['doc_id']}_{chunk['id']}"
                if chunk_identifier in seen_chunk_identifiers:
                    continue

                relevant_chunks_for_llm.append({
                    "id": f"[{citation_id_counter}]",
                    "text": chunk['text'],
                    "original_doc_name": chunk['doc_name']
                })

                citations_for_response.append({
                    "id": str(citation_id_counter),
                    "doc_name": chunk['doc_name'],
                    "text": chunk['text']
                })
                seen_chunk_identifiers.add(chunk_identifier)
                citation_id_counter += 1

            print(
                f"DEBUG RAG: Final Relevant Chunks for LLM (jieba): {[(c['original_doc_name'], c['id']) for c in relevant_chunks_for_llm]}")

        # --- 3. 构建发送给大模型的 Prompt ---
        system_instruction_prefix = "你是一个严谨、专业的AI助手，擅长分析文档并提供清晰、准确、且带有引用的回答。你必须严格依据提供的参考资料进行回答，不允许虚构或依赖你的预训练知识回答参考资料中没有的信息。"

        user_prompt_content = f"用户问题：{user_message_content}"

        if related_file_ids and relevant_chunks_for_llm:
            context_str += "以下是一些用户选择的参考资料。请注意，你的回答必须完全基于这些资料，如果答案不在其中，请明确说明。\n\n"
            for chunk_info in relevant_chunks_for_llm:
                context_str += f"### 参考资料：文档 '{chunk_info['original_doc_name']}'，片段 {chunk_info['id']}\n"
                context_str += "```text\n"
                context_str += chunk_info['text'] + "\n"
                context_str += "```\n\n"

            context_str += "请根据上述提供的参考资料，准确、简洁地回答以下用户的问题。**在回答中，如果使用了参考资料，务必在相关内容后用方括号加数字的形式引用，例如：某个概念的定义[1]，某个论点[2]。** 如果参考资料中没有提及相关信息，请直接回答‘抱歉，根据我当前掌握的资料，无法从提供的参考资料中找到此问题的答案。’\n\n"
        elif related_file_ids and not relevant_chunks_for_llm:
            context_str = "用户提问，但无法从选择的文件中找到与问题相关的参考资料。我将完全根据我已有的知识进行回答。如果您的原始问题是关于特定文档的，请确保文档内容中包含您的问题。请注意，我无法访问您本地的文件。"
        else:
            context_str = "用户提问，未选择任何参考资料。我将完全根据我已有的知识进行回答。\n\n"

        messages_for_llm.append({"role": "system", "content": system_instruction_prefix + context_str})
        messages_for_llm.append({"role": "user", "content": user_prompt_content})

        selected_model = model_settings.get('model', DEFAULT_DEEPSEEK_MODEL)
        if selected_model not in ["deepseek-chat", "deepseek-coder"]:
            print(
                f"Warning: Selected model '{selected_model}' might not be a valid DeepSeek model. Using default: {DEFAULT_DEEPSEEK_MODEL}")
            selected_model = DEFAULT_DEEPSEEK_MODEL

        print(
            f"Calling DeepSeek LLM with model: {selected_model}, temperature: {model_settings.get('temperature', 0.7)}")
        print(f"Prompt sent to LLM:\n{messages_for_llm}")

        completion = openai_client_for_chat_rag.chat.completions.create(
            model=selected_model,
            messages=messages_for_llm,
            temperature=model_settings.get('temperature', 0.7),
            timeout=30.0
        )

        ai_response_content = completion.choices[0].message.content
        print(f"AI Response: {ai_response_content}")

        final_ai_response_json["text"] = ai_response_content
        final_ai_response_json["citations"] = citations_for_response

    except APIConnectionError as e:
        print(f"ERROR: DeepSeek API Connection Error: {e}")
        final_ai_response_json["text"] = f"无法连接到DeepSeek服务，请检查网络或代理设置: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 503
    except RateLimitError as e:
        print(f"ERROR: DeepSeek API Rate Limit Exceeded: {e}")
        final_ai_response_json["text"] = f"请求DeepSeek太频繁，请稍后再试: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 429
    except APIStatusError as e:
        print(
            f"ERROR: DeepSeek API returned non-200 status: {e.status_code} - {e.response.json().get('message', '未知错误')}")
        if e.status_code == 401:
            final_ai_response_json["text"] = f"DeepSeek服务认证失败，请检查API Key是否有效或账户余额。"
            return jsonify({'aiResponse': final_ai_response_json}), 401
        else:
            final_ai_response_json[
                "text"] = f"DeepSeek服务内部错误: {e.status_code} - {e.response.json().get('message', '未知错误')}"
            return jsonify({'aiResponse': final_ai_response_json}), e.status_code
    except Exception as e:
        print(f"ERROR: General error calling DeepSeek API: {e}")
        final_ai_response_json["text"] = f"抱歉，DeepSeek助手在处理您的请求时遇到问题: {str(e)}"
        return jsonify({'aiResponse': final_ai_response_json}), 500

    current_chat_msgs.append({
        'id': f'm_ai_{int(time.time() * 1000) + 1}',
        'sender': 'ai',
        'content': final_ai_response_json['text'],
        'timestamp': datetime.now().strftime('%H:%M'),
        'citations_data': final_ai_response_json['citations']
    })
    mock_chat_messages[new_chat_id] = current_chat_msgs

    if current_chat_entry:
        current_chat_entry['lastActive'] = datetime.now().strftime('%Y-%m-%d %H:%M')
        current_chat_entry['related_file_ids'] = related_file_ids

    save_mock_data()

    response_data = {
        'aiResponse': final_ai_response_json,
        'chatId': new_chat_id
    }
    if not chat_id:
        response_data['newChatId'] = new_chat_id
        response_data['newChatTitle'] = new_chat_title
    return jsonify(response_data), 200


# 7. 导出对话为 DOCX 文件
@app.route('/api/export-chat/<chat_id>', methods=['GET'])
def export_chat(chat_id):
    messages = mock_chat_messages.get(chat_id, [])

    if not messages:
        return jsonify({'error': 'No messages found for this chat ID.'}), 404

    chat_title = "Chat History"
    chat_entry = next((c for c in mock_chat_history if c['id'] == chat_id), None)
    if chat_entry:
        chat_title = chat_entry.get('title', f"Chat-{chat_id}")

    safe_chat_title = secure_filename(chat_title)
    if not safe_chat_title:
        safe_chat_title = f"chat_export_{int(time.time())}"

    document = docx.Document()
    document.add_heading(f"对话历史 - {chat_title}", level=1)

    for msg in messages:
        sender = "您" if msg['sender'] == 'user' else "AI助手"
        timestamp = msg['timestamp']

        p = document.add_paragraph()
        p.add_run(f"[{timestamp}] {sender}: ").bold = True
        p.add_run(msg['content'])

        if msg['sender'] == 'ai' and msg.get('citations_data'):
            document.add_paragraph("--- 引用 ---", style='Intense Quote')
            for citation in msg['citations_data']:
                document.add_paragraph(f"[{citation['id']}] (来自 {citation['doc_name']}): {citation['text']}",
                                       style='List Paragraph')
            document.add_paragraph("-----------")

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{safe_chat_title}_对话历史.docx'
    )


# --- 新功能：答题方法和模板提取相关路由 ---
@app.route('/api/templates/list', methods=['GET'])
def list_templates():
    templates_methods = template_manager.get_all_templates_methods()
    return jsonify(templates_methods)


@app.route('/api/templates/extract-from-file', methods=['POST'])
def extract_templates_from_file():
    data = request.get_json()
    file_id = data.get('fileId')

    if not file_id:
        return jsonify({"error": "File ID is required."}), 400

    file_meta = next((f for f in mock_files_metadata if f['id'] == file_id), None)
    if not file_meta:
        return jsonify({"error": "File not found in metadata."}), 404

    file_path = os.path.join(UPLOAD_FOLDER, file_meta['name'])
    if not os.path.exists(file_path):
        return jsonify({"error": f"File '{file_meta['name']}' not found on disk."}), 404

    full_text_chunks, error_msg = extract_text_from_file(file_path, file_meta['type'])
    if error_msg or not full_text_chunks:
        print(f"DEBUG APP: Failed to extract text chunks from file: {error_msg}")
        return jsonify({"error": f"Failed to extract text from file: {error_msg}"}), 500

    full_text_content = "\n\n".join([c['text'] for c in full_text_chunks])
    file_type_display = file_meta.get('file_type_tag', '未知类型')

    # 1. 使用LLM从整个文档文本中提取Q&A对
    qa_pairs_from_llm = []
    try:
        qa_pairs_from_llm = llm_client_for_tasks.extract_qa_pairs_from_document(
            full_text_content, document_type_tag=file_type_display
        )
        print(
            f"DEBUG APP: LLM extracted Q&A pairs (count: {len(qa_pairs_from_llm)}): {qa_pairs_from_llm[:min(3, len(qa_pairs_from_llm))]}...")  # 打印前3个样本

        if not isinstance(qa_pairs_from_llm, list):
            print(f"DEBUG APP: LLM extract_qa_pairs_from_document returned non-list: {qa_pairs_from_llm}")
            return jsonify(
                {"error": f"LLM failed to extract valid Q&A pairs (returned non-list): {qa_pairs_from_llm}"}), 500
    except Exception as e:
        print(f"ERROR APP: Error calling LLM for Q&A pair extraction: {e}")
        return jsonify({"error": f"Failed to extract Q&A pairs from document: {str(e)}"}), 500

    if not qa_pairs_from_llm:
        print(f"DEBUG APP: No valid Q&A pairs identified by LLM for file {file_meta['name']}.")
        return jsonify({"message": "Successfully analyzed file, but no valid Q&A pairs were identified.",
                        "extracted_data": []}), 200

    # 2. 遍历提取出的Q&A对，并为每个对提取和保存模板/方法
    added_count = 0
    extracted_templates_methods = []

    for i, qa_pair in enumerate(qa_pairs_from_llm):
        question = qa_pair.get("question")
        answer = qa_pair.get("answer")

        if question and answer:
            print(f"DEBUG APP: Processing Q&A pair {i + 1}: Q='{question[:50]}...', A='{answer[:50]}...'")
            try:
                extract_result = question_rewriter_module.extract_and_save_qa_template_method(question, answer)

                print(f"DEBUG APP: Result from question_rewriter_module for pair {i + 1}: {extract_result}")

                if "error" not in extract_result:
                    q_template = extract_result.get("question_template")
                    a_method = extract_result.get("answer_method")
                    if q_template and a_method:
                        if template_manager.add_template_method(q_template, a_method):
                            added_count += 1
                            extracted_templates_methods.append({
                                "question_template": q_template,
                                "answer_method": a_method
                            })
                            print(f"DEBUG APP: Successfully added template/method for pair {i + 1}.")
                        else:
                            print(
                                f"DEBUG APP: Template/Method for pair {i + 1} already exists or failed to add to manager.")
                    else:
                        print(
                            f"WARNING APP: LLM for template extraction returned incomplete data for pair {i + 1}: {extract_result}")
                else:
                    print(
                        f"WARNING APP: Failed to extract template/method for QA pair {i + 1} from LLM: {extract_result['error']}")
            except Exception as e:
                print(f"ERROR APP: Error processing QA pair {i + 1} for template extraction: {e}")
        else:
            print(
                f"WARNING APP: Incomplete Q&A pair returned by LLM (missing question or answer) for pair {i + 1}: {qa_pair}")

    return jsonify({"message": f"Successfully extracted and saved {added_count} templates and methods from file.",
                    "extracted_data": extracted_templates_methods}), 200


@app.route('/api/templates/rewrite-answer', methods=['POST'])
def rewrite_answer_api():
    data = request.get_json()
    question = data.get('question')
    original_answer = data.get('originalAnswer')
    method_index = data.get('methodIndex')

    if not all([question, original_answer, method_index is not None]):
        return jsonify({"error": "Missing question, original answer, or method index."}), 400

    try:
        rewritten_answer = question_rewriter_module.rewrite_answer_with_selected_method(
            question, original_answer, method_index
        )
        if "Error:" in rewritten_answer:
            return jsonify({"error": rewritten_answer}), 500
        return jsonify({"rewrittenAnswer": rewritten_answer}), 200
    except Exception as e:
        print(f"Error in rewrite_answer_api: {e}")
        return jsonify({"error": f"Failed to rewrite answer: {str(e)}"}), 500


# --- 服务器启动 ---
if __name__ == '__main__':
    load_mock_data()
    print(f"\n--- Starting Flask backend server ---")
    print(f"DeepSeek API Key loaded: {'Yes' if deepseek_api_key else 'No'}")
    print(f"DeepSeek Base URL: {DEEPSEEK_BASE_URL}")
    print(f"Uploads directory: {os.path.abspath(UPLOAD_FOLDER)}")
    print(f"Mock data file: {os.path.abspath(MOCK_DATA_FILE)}")
    print(f"Templates/Methods file: {os.path.abspath(TEMPLATES_FILE)}\n")
    app.run(debug=True, port=5000)

